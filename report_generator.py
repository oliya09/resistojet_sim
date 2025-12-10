# report_generator.py — Complete elegant version with enhanced appendix
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime
import matplotlib.pyplot as plt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import os

# ============================================================================
# STYLING CONSTANTS
# ============================================================================
class Colors:
    """Elegant, professional color palette"""
    DARK_GRAY = '#2D3748'
    MEDIUM_GRAY = '#4A5568'
    LIGHT_GRAY = '#718096'
    DARK_BLUE = '#2C5282'
    LIGHT_BLUE = '#4299E1'
    BG_LIGHT = '#F7FAFC'
    BG_SOFT = '#EDF2F7'
    BORDER = '#E2E8F0'

# ============================================================================
# DOCX REPORT GENERATOR
# ============================================================================

def generate_docx_report(sim_df, metrics, inputs, plot_paths, out_path, 
                        include_methodology=False, include_raw_data=False,
                        include_recommendations=False, report_notes="", user_name="User",
                        use_extended_version=False, chamber_heater_on=False, tank_heater_on=False,
                        use_regulator=False):
    """Generate DOCX report with elegant appendix"""
    
    doc = Document()
    
    # Set default styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ==================== TITLE SECTION ====================
    # Title
    title = doc.add_heading('Resistojet Simulation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtle divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 40)
    run.font.color.rgb = RGBColor(220, 220, 220)
    run.font.size = Pt(8)
    
    # User info and date
    doc.add_paragraph(f"Generated for: {user_name}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # ==================== PRESETS SECTION ====================
    doc.add_heading('PRESETS', level=1)
    
    # Create presets table
    presets_table = doc.add_table(rows=1, cols=2)
    presets_table.style = 'Light Grid Accent 1'
    
    # Header with subtle styling
    hdr_cells = presets_table.rows[0].cells
    hdr_cells[0].text = "Parameter"
    hdr_cells[1].text = "Value"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Basic presets
    basic_presets = [
        ("Propellant", inputs.get('Propellant', 'N/A')),
        ("Throat diameter Dt [mm]", inputs.get('Throat Diameter Dt [mm]', 'N/A')),
        ("Exit diameter De [mm]", inputs.get('Exit Diameter De [mm]', 'N/A')),
        ("Chamber material", inputs.get('Chamber Material', 'N/A')),
        ("Tank initial temperature Tt [K]", inputs.get('Initial Tank Temperature [K]', 'N/A')),
        ("Initial saturated pressure (est.) Pt [bar]", inputs.get('Saturated Pressure [bar]', 'N/A')),  # Changed from 'Initial Tank Pressure [bar]' to 'Saturated Pressure [bar]'
        ("Propellant mass [kg]", inputs.get('Propellant Mass [kg]', 'N/A')),
        ("Volume [L]", inputs.get('Propellant Volume [L]', 'N/A')),
        ("Ambient/back pressure [Pa]", inputs.get('Ambient/Back Pressure [Pa]', 'N/A')),
        ("Timestep dt [s]", inputs.get('Timestep dt [s]', 'N/A')),
        ("Simulation time [s]", inputs.get('Total Simulation Time [s]', 'N/A'))
    ]
    
    # Add basic presets
    for preset_name, preset_value in basic_presets:
        row_cells = presets_table.add_row().cells
        row_cells[0].text = preset_name
        row_cells[1].text = str(preset_value)
    
    # Extended version presets
    if use_extended_version:
        # Add separator
        row_cells = presets_table.add_row().cells
        row_cells[0].text = "─" * 20
        row_cells[1].text = "─" * 20
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(200, 200, 200)
        
        extended_presets = [
            ("Enable Chamber Heater", "Yes" if chamber_heater_on else "No"),
            ("Chamber Heater Material", inputs.get('Chamber Heater Material', 'N/A') if chamber_heater_on else "N/A"),
            ("Chamber Heater Power [W]", inputs.get('Chamber Heater Power [W]', 'N/A') if chamber_heater_on else "N/A"),
            ("Chamber Heater Efficiency [%]", inputs.get('Chamber Heater Efficiency [%]', 'N/A') if chamber_heater_on else "N/A"),
            ("Chamber Heater Surface Area [cm²]", inputs.get('Chamber Heater Surface Area [cm²]', 'N/A') if chamber_heater_on else "N/A"),
            ("Enable Tank Heater", "Yes" if tank_heater_on else "No"),
            ("Tank Heater Power [W]", inputs.get('Tank Heater Power [W]', 'N/A') if tank_heater_on else "N/A"),
            ("Tank Heater Efficiency [%]", inputs.get('Tank Heater Efficiency [%]', 'N/A') if tank_heater_on else "N/A"),
            ("Enable Pressure Regulator", "Yes" if use_regulator else "No"),
            ("Regulator Setpoint Pressure [bar]", inputs.get('Regulator Setpoint [bar]', 'N/A') if use_regulator else "N/A")
        ]
        
        for preset_name, preset_value in extended_presets:
            row_cells = presets_table.add_row().cells
            row_cells[0].text = preset_name
            row_cells[1].text = str(preset_value)
    
    doc.add_paragraph()
    
    # ==================== PERFORMANCE METRICS ====================
    doc.add_heading('Performance Metrics', level=1)
    add_table_to_docx(doc, "Metric", "Value", metrics)
    
    doc.add_paragraph()
    
    # ==================== PLOTS SECTION ====================
    for plot_title, plot_path in plot_paths.items():
        if os.path.exists(plot_path):
            doc.add_picture(plot_path, width=Inches(6))
            # Add caption
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run(plot_title)
            caption_run.font.size = Pt(9)
            caption_run.font.color.rgb = RGBColor(100, 100, 100)
            caption_run.italic = True
            doc.add_paragraph()
    
    # ==================== ADDITIONAL NOTES ====================
    if report_notes:
        doc.add_heading('Additional Notes', level=1)
        note_para = doc.add_paragraph()
        note_para.paragraph_format.left_indent = Inches(0.25)
        note_run = note_para.add_run(report_notes)
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = RGBColor(80, 80, 80)
    
    # ==================== APPENDIX: METHODOLOGY ====================
    if include_methodology:
        doc.add_page_break()
        
        # Appendix Header
        doc.add_heading('Appendix A: Resistojet Thrust Model', level=1)
        
        # Add introductory paragraph
        intro = doc.add_paragraph()
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        intro_run = intro.add_run(
            'This appendix documents the mathematical models and assumptions used in the '
            'resistojet simulation. The implementation follows established thermodynamic '
            'principles and conservation laws.'
        )
        intro_run.italic = True
        intro_run.font.color.rgb = RGBColor(80, 80, 80)
        
        doc.add_paragraph()
        
        # ========== A.1 SYSTEM OVERVIEW ==========
        doc.add_heading('A.1 System Overview', level=2)
        
        overview = doc.add_paragraph('The resistojet is modeled as a system of three main components:')
        overview.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Components list with subtle styling
        components = [
            ('Propellant Tank', 'Stores liquid propellant in vapor-liquid equilibrium'),
            ('Heating Chamber', 'Electrically heats the incoming vapor'),
            ('Nozzle', 'Accelerates the hot gas through converging-diverging geometry')
        ]
        
        for i, (name, desc) in enumerate(components, 1):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
            run_name = p.add_run(f"{name}: ")
            run_name.bold = True
            run_desc = p.add_run(desc)
            run_desc.font.color.rgb = RGBColor(60, 60, 60)
        
        summary = doc.add_paragraph()
        summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        summary.add_run('The model solves coupled conservation equations for mass and energy across these components.')
        
        doc.add_paragraph()
        
        # ========== A.2 THERMODYNAMIC FOUNDATION ==========
        doc.add_heading('A.2 Thermodynamic Foundation', level=2)
        
        doc.add_heading('A.2.1 Temperature-Dependent Properties', level=3)
        
        doc.add_paragraph('Specific heat capacity (at constant pressure) is modeled using a polynomial fit from NASA CEA data:')
        
        # Equation with better formatting
        eq1 = doc.add_paragraph()
        eq1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq1.paragraph_format.space_before = Pt(8)
        eq1.paragraph_format.space_after = Pt(8)
        eq1_run = eq1.add_run('cₚ(T) = a + bT + cT² + dT³')
        eq1_run.font.name = 'Cambria Math'
        eq1_run.font.size = Pt(11)
        eq1_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        
        doc.add_paragraph('where a, b, c, d are substance-specific coefficients.')
        
        doc.add_paragraph()
        doc.add_paragraph('The specific heat ratio is computed as:')
        
        eq2 = doc.add_paragraph()
        eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq2.paragraph_format.space_before = Pt(8)
        eq2.paragraph_format.space_after = Pt(12)
        eq2_run = eq2.add_run('γ(T) = cₚ(T) / [cₚ(T) - R]')
        eq2_run.font.name = 'Cambria Math'
        eq2_run.font.size = Pt(11)
        eq2_run.font.color.rgb = RGBColor(0, 0, 139)
        
        doc.add_paragraph('where R = Rᵤ/M is the specific gas constant.')
        
        doc.add_paragraph()
        
        # ========== A.3 PROPELLANT TANK MODEL ==========
        doc.add_heading('A.3 Propellant Tank Model', level=2)
        
        doc.add_heading('A.3.1 Vapor-Liquid Equilibrium', level=3)
        
        doc.add_paragraph('Saturation pressure calculated using the Wagner equation:')
        
        eq3 = doc.add_paragraph()
        eq3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq3.paragraph_format.space_before = Pt(8)
        eq3.paragraph_format.space_after = Pt(8)
        eq3_run = eq3.add_run('ln(P/Pc) = [A(1-Tᵣ) + B(1-Tᵣ)¹·⁵ + C(1-Tᵣ)³ + D(1-Tᵣ)⁶] / Tᵣ')
        eq3_run.font.name = 'Cambria Math'
        eq3_run.font.size = Pt(11)
        eq3_run.font.color.rgb = RGBColor(0, 0, 139)
        
        doc.add_paragraph('where Pc is critical pressure, Tᵣ = T/Tc is reduced temperature, and A,B,C,D are fluid-specific Wagner coefficients.')
        
        doc.add_paragraph()
        
        doc.add_heading('A.3.2 Tank Temperature Evolution', level=3)
        
        doc.add_paragraph('Liquid temperature changes due to external heating and evaporative cooling:')
        
        # Temperature equations with clear formatting
        temp_changes = [
            ('1. Heating from external power:', 'ΔT_heat = (Q_tank η_tank Δt) / (m cₚ,liq)'),
            ('2. Evaporative cooling from mass outflow:', 'ΔT_cool = -( m_dot Lᵥ Δt) / (m cₚ,liq)'),
            ('3. Net temperature change:', 'T_new = T_current + ΔT_heat + ΔT_cool')
        ]
        
        for label, equation in temp_changes:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            label_run = p.add_run(label + ' ')
            label_run.bold = True
            eq_run = p.add_run(equation)
            eq_run.font.name = 'Cambria Math'
            eq_run.font.size = Pt(10)
            eq_run.font.color.rgb = RGBColor(60, 60, 160)
        
        doc.add_paragraph()
        
        # ========== A.4 HEATING CHAMBER MODEL ==========
        doc.add_heading('A.4 Heating Chamber Model', level=2)
        
        doc.add_heading('A.4.1 Energy Balance', level=3)
        
        doc.add_paragraph('Steady-state energy balance:')
        
        eq4 = doc.add_paragraph()
        eq4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq4.paragraph_format.space_before = Pt(8)
        eq4.paragraph_format.space_after = Pt(8)
        eq4_run = eq4.add_run(' m_dot cₚ(T_c) (T_c - T_t) = Q_chamber η_chamber - Q_loss')
        eq4_run.font.name = 'Cambria Math'
        eq4_run.font.size = Pt(11)
        eq4_run.font.color.rgb = RGBColor(0, 0, 139)
        
        doc.add_paragraph()
        doc.add_paragraph('Radiative heat loss:')
        
        eq5 = doc.add_paragraph()
        eq5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq5.paragraph_format.space_before = Pt(8)
        eq5.paragraph_format.space_after = Pt(12)
        eq5_run = eq5.add_run('Q_loss = σ ε A_h (T_c⁴ - T_amb⁴)')
        eq5_run.font.name = 'Cambria Math'
        eq5_run.font.size = Pt(11)
        eq5_run.font.color.rgb = RGBColor(0, 0, 139)
        
        doc.add_heading('A.4.2 Temperature Limit', level=3)
        
        doc.add_paragraph('Chamber temperature constraint:')
        
        eq6 = doc.add_paragraph()
        eq6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq6.paragraph_format.space_before = Pt(8)
        eq6.paragraph_format.space_after = Pt(8)
        eq6_run = eq6.add_run('T_c ≤ min( (Q_chamber/(σ ε A_h))¹/⁴, T_max,mat )')
        eq6_run.font.name = 'Cambria Math'
        eq6_run.font.size = Pt(11)
        eq6_run.font.color.rgb = RGBColor(0, 0, 139)
        
        doc.add_paragraph('where T_max,mat is the maximum allowable material temperature.')
        
        doc.add_heading('A.4.3 Numerical Solution', level=3)
        
        doc.add_paragraph('The nonlinear energy balance equation is solved for T_c using the Newton-Raphson iterative method.')
        
        doc.add_paragraph()
        
        # ========== A.5 NOZZLE FLOW MODEL ==========
        doc.add_heading('A.5 Nozzle Flow Model', level=2)
        
        doc.add_heading('A.5.1 Throat Conditions', level=3)
        
        doc.add_paragraph('Choked mass flow rate:')
        
        eq7 = doc.add_paragraph()
        eq7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq7.paragraph_format.space_before = Pt(8)
        eq7.paragraph_format.space_after = Pt(8)
        eq7_run = eq7.add_run('m_dot = A_t P_c √[γ/(R T_c)] [(γ+1)/2]^[-(γ+1)/(2(γ-1))]')
        eq7_run.font.name = 'Cambria Math'
        eq7_run.font.size = Pt(11)
        eq7_run.font.color.rgb = RGBColor(0, 0, 139)
        
        doc.add_paragraph('where A_t is throat area and P_c is chamber pressure.')
        
        doc.add_heading('A.5.2 Exit Mach Number', level=3)
        
        doc.add_paragraph('Exit Mach number M_e from area ratio:')
        
        eq8 = doc.add_paragraph()
        eq8.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq8.paragraph_format.space_before = Pt(8)
        eq8.paragraph_format.space_after = Pt(8)
        eq8_run = eq8.add_run('A_e/A_t = (1/M_e)[2/(γ+1) (1 + (γ-1)/2 M_e²)]^[(γ+1)/(2(γ-1))]')
        eq8_run.font.name = 'Cambria Math'
        eq8_run.font.size = Pt(11)
        eq8_run.font.color.rgb = RGBColor(0, 0, 139)
        
        doc.add_heading('A.5.3 Exit Flow Properties', level=3)
        
        doc.add_paragraph('Assuming isentropic expansion:')
        
        isentropic_eqns = [
            'T_e = T_c / [1 + (γ-1)/2 M_e²]',
            'V_e = M_e √(γ R T_e)',
            'p_e = P_c (T_e/T_c)^[γ/(γ-1)]'
        ]
        
        for eqn in isentropic_eqns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            eqn_run = p.add_run(eqn)
            eqn_run.font.name = 'Cambria Math'
            eqn_run.font.size = Pt(10)
            eqn_run.font.color.rgb = RGBColor(60, 60, 160)
        
        doc.add_paragraph()
        
        # ========== A.6 THRUST AND PERFORMANCE METRICS ==========
        doc.add_heading('A.6 Thrust and Performance Metrics', level=2)
        
        doc.add_heading('A.6.1 Thrust', level=3)
        
        doc.add_paragraph('F = m_dot V_e + (p_e - p_amb) A_e')
        doc.add_paragraph('where p_amb is ambient pressure.')
        
        doc.add_heading('A.6.2 Specific Impulse', level=3)
        
        eq9 = doc.add_paragraph()
        eq9.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq9.paragraph_format.space_before = Pt(8)
        eq9.paragraph_format.space_after = Pt(8)
        eq9_run = eq9.add_run('I_sp = F / (m_dot g₀)')
        eq9_run.font.name = 'Cambria Math'
        eq9_run.font.size = Pt(11)
        eq9_run.font.color.rgb = RGBColor(0, 0, 139)
        
        doc.add_paragraph('where g₀ = 9.80665 m/s².')
        
        doc.add_heading('A.6.3 Total Impulse', level=3)
        
        doc.add_paragraph('For time-stepped simulation:')
        
        eq10 = doc.add_paragraph()
        eq10.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq10.paragraph_format.space_before = Pt(8)
        eq10.paragraph_format.space_after = Pt(8)
        eq10_run = eq10.add_run('I_total = Σ F Δt')
        eq10_run.font.name = 'Cambria Math'
        eq10_run.font.size = Pt(11)
        eq10_run.font.color.rgb = RGBColor(0, 0, 139)
        
        doc.add_heading('A.6.4 Thrust Efficiency', level=3)
        
        eq11 = doc.add_paragraph()
        eq11.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq11.paragraph_format.space_before = Pt(8)
        eq11.paragraph_format.space_after = Pt(12)
        eq11_run = eq11.add_run('η_thruster = F² / [2 m_dot(Q_tank + Q_chamber)]')
        eq11_run.font.name = 'Cambria Math'
        eq11_run.font.size = Pt(11)
        eq11_run.font.color.rgb = RGBColor(0, 0, 139)
        
        doc.add_paragraph()
        
        # ========== A.7 KEY ASSUMPTIONS ==========
        doc.add_heading('A.7 Key Assumptions', level=2)
        
        assumptions = [
            'Vapor behaves as a perfect (ideal) gas in chamber and nozzle',
            'Nozzle flow is isentropic (no friction, shocks, or boundary layer losses)',
            'Flow is choked at the throat',
            'Chamber temperature and pressure are uniform',
            'No chemical reactions or composition changes occur',
            'Vaporization at tank outlet occurs instantaneously at saturation conditions'
        ]
        
        for i, assumption in enumerate(assumptions, 1):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(assumption)
        
        doc.add_paragraph()
        
        # ========== A.8 DATA SOURCES ==========
        doc.add_heading('A.8 Data Sources', level=2)
        
        sources = [
            ('• Fluid properties:', 'NIST Chemistry WebBook'),
            ('• Wagner equation coefficients:', 'DIPPR Project 801'),
            ('• High-fidelity thermodynamic data:', 'REFPROP database'),
            ('• Fluid cₚ(T) polynomial coefficients:', 'NASA Chemical Equilibrium with Applications (CEA)')
        ]
        
        for label, source in sources:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            label_run = p.add_run(label + ' ')
            label_run.bold = True
            source_run = p.add_run(source)
            source_run.font.color.rgb = RGBColor(80, 80, 80)
        
        # Add closing divider
        doc.add_paragraph()
        closing_p = doc.add_paragraph()
        closing_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        closing_run = closing_p.add_run('─ End of Appendix ─')
        closing_run.font.color.rgb = RGBColor(180, 180, 180)
        closing_run.font.size = Pt(9)
        closing_run.italic = True
    
    # ==================== FOOTER ====================
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Resistojet Simulation Report - Generated for {user_name} - {datetime.now().strftime('%Y-%m-%d')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(out_path)

# ============================================================================
# PDF REPORT GENERATOR
# ============================================================================

def generate_pdf_report(sim_df, metrics, inputs, plot_paths, out_path,
                       include_methodology=False, include_raw_data=False,
                       include_recommendations=False, report_notes="", user_name="User",
                       use_extended_version=False, chamber_heater_on=False, tank_heater_on=False,
                       use_regulator=False):
    
    doc = SimpleDocTemplate(out_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # ==================== CUSTOM STYLES ====================
    # Title style
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Title'],
        fontSize=22,
        textColor=colors.HexColor('#000000'),
        alignment=1,
        spaceAfter=24,
        fontName='Helvetica-Bold'
    )
    
    # Heading styles
    heading_style = ParagraphStyle(
        'Heading1',
        parent=styles['Heading1'],
        fontSize=14,
        textColor=colors.HexColor('#222222'),
        spaceBefore=20,
        spaceAfter=10,
        fontName='Helvetica-Bold'
    )
    
    subheading_style = ParagraphStyle(
        'Heading2',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.HexColor('#333333'),
        spaceBefore=15,
        spaceAfter=8,
        fontName='Helvetica-Bold'
    )
    
    small_subheading_style = ParagraphStyle(
        'Heading3',
        parent=styles['Heading2'],
        fontSize=11,
        textColor=colors.HexColor('#444444'),
        spaceBefore=12,
        spaceAfter=6,
        fontName='Helvetica-Bold'
    )
    
    # Normal text
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#000000'),
        spaceAfter=6,
        fontName='Helvetica'
    )
    
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#000000'),
        leftIndent=20,
        spaceAfter=4
    )
    
    equation_style = ParagraphStyle(
        'Equation',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#000000'),  # Dark blue
        alignment=1,
        fontName='Courier-Bold',
        spaceBefore=10,
        spaceAfter=10
    )
    
    caption_style = ParagraphStyle(
        'Caption',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#666666'),
        alignment=1,
        fontName='Helvetica-Oblique',
        spaceBefore=4,
        spaceAfter=12
    )
    
    # ==================== TITLE SECTION ====================
    story.append(Paragraph("Resistojet Simulation Report", title_style))
    story.append(Spacer(1, 0.1*inch))
    
    # Divider line
    story.append(HRFlowable(width="80%", thickness=0.5, color=colors.HexColor('#CCCCCC')))
    story.append(Spacer(1, 0.15*inch))
    
    # User info
    story.append(Paragraph(f"<b>Generated for:</b> {user_name}", normal_style))
    story.append(Paragraph(f"<b>Date:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
    story.append(Spacer(1, 0.1*inch))
    
    # ==================== PRESETS SECTION ====================
    story.append(Paragraph("PRESETS", heading_style))
    story.append(Spacer(1, 0.05*inch))
    
    # Prepare presets data
    presets_data = [["Parameter", "Value"]]
    
    # Basic presets
    basic_presets = [
        ("Propellant", inputs.get('Propellant', 'N/A')),
        ("Throat diameter Dt [mm]", inputs.get('Throat Diameter Dt [mm]', 'N/A')),
        ("Exit diameter De [mm]", inputs.get('Exit Diameter De [mm]', 'N/A')),
        ("Chamber material", inputs.get('Chamber Material', 'N/A')),
        ("Tank initial temperature Tt [K]", inputs.get('Initial Tank Temperature [K]', 'N/A')),
        ("Initial saturated pressure (est.) Pt [bar]", inputs.get('Saturated Pressure [bar]', 'N/A')),
        ("Propellant mass [kg]", inputs.get('Propellant Mass [kg]', 'N/A')),
        ("Volume [L]", inputs.get('Propellant Volume [L]', 'N/A')),
        ("Ambient/back pressure [Pa]", inputs.get('Ambient/Back Pressure [Pa]', 'N/A')),
        ("Timestep dt [s]", inputs.get('Timestep dt [s]', 'N/A')),
        ("Simulation time [s]", inputs.get('Total Simulation Time [s]', 'N/A'))
    ]
    
    for preset_name, preset_value in basic_presets:
        presets_data.append([preset_name, str(preset_value)])
    
    # Extended version presets
    if use_extended_version:
        presets_data.append(["", ""])  # Separator
        
        extended_presets = [
            ("Enable Chamber Heater", "Yes" if chamber_heater_on else "No"),
            ("Chamber Heater Material", inputs.get('Chamber Heater Material', 'N/A') if chamber_heater_on else "N/A"),
            ("Chamber Heater Power [W]", inputs.get('Chamber Heater Power [W]', 'N/A') if chamber_heater_on else "N/A"),
            ("Chamber Heater Efficiency [%]", inputs.get('Chamber Heater Efficiency [%]', 'N/A') if chamber_heater_on else "N/A"),
            ("Chamber Heater Surface Area [cm²]", inputs.get('Chamber Heater Surface Area [cm²]', 'N/A') if chamber_heater_on else "N/A"),
            ("Enable Tank Heater", "Yes" if tank_heater_on else "No"),
            ("Tank Heater Power [W]", inputs.get('Tank Heater Power [W]', 'N/A') if tank_heater_on else "N/A"),
            ("Tank Heater Efficiency [%]", inputs.get('Tank Heater Efficiency [%]', 'N/A') if tank_heater_on else "N/A"),
            ("Enable Pressure Regulator", "Yes" if use_regulator else "No"),
            ("Regulator Setpoint Pressure [bar]", inputs.get('Regulator Setpoint [bar]', 'N/A') if use_regulator else "N/A")
        ]
        
        for preset_name, preset_value in extended_presets:
            presets_data.append([preset_name, str(preset_value)])
    
    # Create presets table
    presets_table = Table(presets_data, colWidths=[3.5*inch, 2.5*inch])
    presets_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F0F0F0')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D0D0D0')),
    ]))
    story.append(presets_table)
    story.append(Spacer(1, 0.15*inch))
    
    # ==================== PERFORMANCE METRICS ====================
    story.append(Paragraph("Performance Metrics", heading_style))
    story.append(Spacer(1, 0.05*inch))
    
    metric_data = [["Metric", "Value"]]
    for key, value in metrics.items():
        metric_data.append([key, str(value)])
    
    metric_table = Table(metric_data, colWidths=[3.5*inch, 2.5*inch])
    metric_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F0F0F0')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D0D0D0')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8F8F8')]),
    ]))
    story.append(metric_table)
    story.append(Spacer(1, 0.2*inch))
    
    # ==================== PLOTS SECTION ====================
    for plot_title, plot_path in plot_paths.items():
        if os.path.exists(plot_path):
            img = Image(plot_path, width=6*inch, height=4*inch)
            story.append(img)
            story.append(Paragraph(plot_title, caption_style))
    
    # ==================== ADDITIONAL NOTES ====================
    if report_notes:
        story.append(Paragraph("Additional Notes", heading_style))
        story.append(Paragraph(report_notes, normal_style))
        story.append(Spacer(1, 0.1*inch))
    
    # ==================== APPENDIX: METHODOLOGY ====================
    if include_methodology:
        story.append(PageBreak())
        story.append(Spacer(1, 0.3*inch))
        
        # Appendix Title
        story.append(Paragraph("Appendix A: Resistojet Thrust Model", title_style))
        story.append(Spacer(1, 0.05*inch))
        
        # Introduction
        intro_text = """
        <i>This appendix details the mathematical models and assumptions used in the 
        resistojet simulation. The implementation follows established thermodynamic 
        principles and conservation laws.</i>
        """
        story.append(Paragraph(intro_text, normal_style))
        story.append(Spacer(1, 0.2*inch))
        
        # ========== A.1 SYSTEM OVERVIEW ==========
        story.append(Paragraph("A.1 System Overview", heading_style))
        story.append(Paragraph("The resistojet is modeled as a system of three main components:", normal_style))
        
        components = [
            ("<b>Propellant Tank:</b> Stores liquid propellant in vapor-liquid equilibrium"),
            ("<b>Heating Chamber:</b> Electrically heats the incoming vapor"),
            ("<b>Nozzle:</b> Accelerates the hot gas through converging-diverging geometry")
        ]
        
        for component in components:
            story.append(Paragraph(component, bullet_style))
        
        story.append(Paragraph("The model solves coupled conservation equations for mass and energy across these components.", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        # ========== A.2 THERMODYNAMIC FOUNDATION ==========
        story.append(Paragraph("A.2 Thermodynamic Foundation", heading_style))
        story.append(Paragraph("A.2.1 Temperature-Dependent Properties", subheading_style))
        
        story.append(Paragraph("Specific heat capacity (at constant pressure) is modeled using a polynomial fit from NASA CEA data:", normal_style))
        
        eq1 = "c<sub>p</sub>(T) = a + bT + cT<sup>2</sup> + dT<sup>3</sup>"
        story.append(Paragraph(eq1, equation_style))
        story.append(Paragraph("where a, b, c, d are substance-specific coefficients.", normal_style))
        story.append(Spacer(1, 0.05*inch))
        
        story.append(Paragraph("The specific heat ratio is computed as:", normal_style))
        
        eq2 = "γ(T) = c<sub>p</sub>(T) / [c<sub>p</sub>(T) - R]"
        story.append(Paragraph(eq2, equation_style))
        story.append(Paragraph("where R = R<sub>u</sub>/M is the specific gas constant.", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        # ========== A.3 PROPELLANT TANK MODEL ==========
        story.append(Paragraph("A.3 Propellant Tank Model", heading_style))
        story.append(Paragraph("A.3.1 Vapor-Liquid Equilibrium", subheading_style))
        
        story.append(Paragraph("Saturation pressure calculated using the Wagner equation:", normal_style))
        
        eq3 = "ln(P/P<sub>c</sub>) = [A(1-T<sub>r</sub>) + B(1-T<sub>r</sub>)<sup>1.5</sup> + C(1-T<sub>r</sub>)<sup>3</sup> + D(1-T<sub>r</sub>)<sup>6</sup>] / T<sub>r</sub>"
        story.append(Paragraph(eq3, equation_style))
        story.append(Paragraph("where P<sub>c</sub> is critical pressure, T<sub>r</sub> = T/T<sub>c</sub> is reduced temperature,", normal_style))
        story.append(Paragraph("and A,B,C,D are fluid-specific Wagner coefficients.", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph("A.3.2 Tank Temperature Evolution", subheading_style))
        story.append(Paragraph("Liquid temperature changes due to external heating and evaporative cooling:", normal_style))
        
        temp_changes = [
            ("1. <b>Heating from external power:</b>", "ΔT<sub>heat</sub> = (Q<sub>tank</sub> η<sub>tank</sub> Δt) / (m c<sub>p,liq</sub>)"),
            ("2. <b>Evaporative cooling from mass outflow:</b>", "ΔT<sub>cool</sub> = -(m_dot L<sub>v</sub> Δt) / (m c<sub>p,liq</sub>)"),
            ("3. <b>Net temperature change:</b>", "T<sub>new</sub> = T<sub>current</sub> + ΔT<sub>heat</sub> + ΔT<sub>cool</sub>")
        ]
        
        for label, eqn in temp_changes:
            story.append(Paragraph(label, bullet_style))
            p = Paragraph(eqn, ParagraphStyle('EquationInline', parent=normal_style, 
                                             fontName='Courier', fontSize=10,
                                             leftIndent=40))
            story.append(p)
        
        story.append(Spacer(1, 0.15*inch))
        
        # ========== A.4 HEATING CHAMBER MODEL ==========
        story.append(Paragraph("A.4 Heating Chamber Model", heading_style))
        story.append(Paragraph("A.4.1 Energy Balance", subheading_style))
        
        story.append(Paragraph("Steady-state energy balance:", normal_style))
        
        eq4 = "m_dot c<sub>p</sub>(T<sub>c</sub>) (T<sub>c</sub> - T<sub>t</sub>) = Q<sub>chamber</sub> η<sub>chamber</sub> - Q<sub>loss</sub>"
        story.append(Paragraph(eq4, equation_style))
        story.append(Spacer(1, 0.05*inch))
        
        story.append(Paragraph("Radiative heat loss:", normal_style))
        
        eq5 = "Q<sub>loss</sub> = σ ε A<sub>h</sub> (T<sub>c</sub><sup>4</sup> - T<sub>amb</sub><sup>4</sup>)"
        story.append(Paragraph(eq5, equation_style))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph("A.4.2 Temperature Limit", subheading_style))
        story.append(Paragraph("Chamber temperature constraint:", normal_style))
        
        eq6 = "T<sub>c</sub> ≤ min( (Q<sub>chamber</sub>/(σ ε A<sub>h</sub>))<sup>1/4</sup>, T<sub>max,mat</sub> )"
        story.append(Paragraph(eq6, equation_style))
        story.append(Paragraph("where T<sub>max,mat</sub> is the maximum allowable material temperature.", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph("A.4.3 Numerical Solution", subheading_style))
        story.append(Paragraph("The nonlinear energy balance equation is solved for T<sub>c</sub> using the Newton-Raphson iterative method.", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        # ========== A.5 NOZZLE FLOW MODEL ==========
        story.append(Paragraph("A.5 Nozzle Flow Model", heading_style))
        story.append(Paragraph("A.5.1 Throat Conditions", subheading_style))
        
        story.append(Paragraph("Choked mass flow rate:", normal_style))
        
        eq7 = "m_dot = A<sub>t</sub> P<sub>c</sub> √[γ/(R T<sub>c</sub>)] [(γ+1)/2]<sup>-(γ+1)/(2(γ-1))</sup>"
        story.append(Paragraph(eq7, equation_style))
        story.append(Paragraph("where A<sub>t</sub> is throat area and P<sub>c</sub> is chamber pressure.", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph("A.5.2 Exit Mach Number", subheading_style))
        story.append(Paragraph("Exit Mach number M<sub>e</sub> from area ratio:", normal_style))
        
        eq8 = "A<sub>e</sub>/A<sub>t</sub> = (1/M<sub>e</sub>)[2/(γ+1) (1 + (γ-1)/2 M<sub>e</sub><sup>2</sup>)]<sup>(γ+1)/(2(γ-1))</sup>"
        story.append(Paragraph(eq8, equation_style))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph("A.5.3 Exit Flow Properties", subheading_style))
        story.append(Paragraph("Assuming isentropic expansion:", normal_style))
        
        isentropic_eqns = [
            "T<sub>e</sub> = T<sub>c</sub> / [1 + (γ-1)/2 M<sub>e</sub><sup>2</sup>]",
            "V<sub>e</sub> = M<sub>e</sub> √(γ R T<sub>e</sub>)",
            "p<sub>e</sub> = P<sub>c</sub> (T<sub>e</sub>/T<sub>c</sub>)<sup>γ/(γ-1)</sup>"
        ]
        
        for eqn in isentropic_eqns:
            story.append(Paragraph(eqn, bullet_style))
        
        story.append(Spacer(1, 0.15*inch))
        
        # ========== A.6 THRUST AND PERFORMANCE METRICS ==========
        story.append(Paragraph("A.6 Thrust and Performance Metrics", heading_style))
        story.append(Paragraph("A.6.1 Thrust", subheading_style))
        
        eq9 = "F = m_dot V<sub>e</sub> + (p<sub>e</sub> - p<sub>amb</sub>) A<sub>e</sub>"
        story.append(Paragraph(eq9, equation_style))
        story.append(Paragraph("where p<sub>amb</sub> is ambient pressure.", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph("A.6.2 Specific Impulse", subheading_style))
        
        eq10 = "I<sub>sp</sub> = F / (m_dot g<sub>0</sub>)"
        story.append(Paragraph(eq10, equation_style))
        story.append(Paragraph("where g<sub>0</sub> = 9.80665 m/s<sup>2</sup>.", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph("A.6.3 Total Impulse", subheading_style))
        story.append(Paragraph("For time-stepped simulation:", normal_style))
        
        eq11 = "I<sub>total</sub> = Σ F Δt"
        story.append(Paragraph(eq11, equation_style))
        story.append(Spacer(1, 0.15*inch))
        
        # ========== A.7 KEY ASSUMPTIONS ==========
        story.append(Paragraph("A.7 Key Assumptions", heading_style))
        
        assumptions = [
            "Vapor behaves as a perfect (ideal) gas in chamber and nozzle",
            "Nozzle flow is isentropic (no friction, shocks, or boundary layer losses)",
            "Flow is choked at the throat",
            "Chamber temperature and pressure are uniform",
            "No chemical reactions or composition changes occur",
            "Vaporization at tank outlet occurs instantaneously at saturation conditions"
        ]
        
        for i, assumption in enumerate(assumptions, 1):
            story.append(Paragraph(f"{i}. {assumption}", bullet_style))
        
        story.append(Spacer(1, 0.15*inch))
        
        # ========== A.8 DATA SOURCES ==========
        story.append(Paragraph("A.8 Data Sources", heading_style))
        
        sources = [
            ("• <b>Fluid c<sub>p</sub>(T) polynomial coefficients:</b>", "NASA Chemical Equilibrium with Applications (CEA)"),
            ("• <b>Wagner equation coefficients:</b>", "DIPPR Project 801"),
            ("• <b>High-fidelity thermodynamic data:</b>", "REFPROP database"),
            ("• <b>Fluid properties:</b>", "NIST Chemistry WebBook")
        ]
        
        for label, source in sources:
            story.append(Paragraph(f"{label} {source}", bullet_style))
        
        # Closing
        story.append(Spacer(1, 0.3*inch))
        story.append(HRFlowable(width="40%", thickness=0.5, color=colors.HexColor('#CCCCCC')))
        story.append(Paragraph("<i>End of Appendix</i>", caption_style))
    
    doc.build(story)

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def add_table_to_docx(doc, col1_header, col2_header, data_dict):
    """Helper to add a table to DOCX document"""
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = col1_header
    hdr_cells[1].text = col2_header
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    for key, value in data_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)

def add_dataframe_to_docx(doc, df):
    '''Helper to add DataFrame to DOCX document'''
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        # Make header bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
